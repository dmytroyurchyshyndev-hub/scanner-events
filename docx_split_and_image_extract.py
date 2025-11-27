#!/usr/bin/env python3
"""
Structured DOCX Image Extractor & Page Splitter

Workflow:
1. Load DOCX
2. Extract images and record positions
3. Create cleaned DOCX without images
4. Convert cleaned DOCX to PDF
5. Split PDF into pages
6. Save metadata JSON with image info and page files

Requirements:
pip install python-docx docx2pdf pymupdf tqdm
"""

import argparse
import os
import uuid
import json
import shutil
import subprocess
from pathlib import Path
from tqdm import tqdm
from docx import Document
import re
import fitz


NS_RID_PATTERN = re.compile(r'r:embed\"(rId[0-9]+)\"')


class DocxProcessor:
    def __init__(self, input_path: Path, outdir: Path, attempt_hash: str = None):
        self.input_path = input_path
        self.hash = attempt_hash or uuid.uuid4().hex[:12]
        self.base_out = outdir / self.hash
        self.images_dir = self.base_out / "images"
        self.pages_dir = self.base_out / "pages"
        self.work_dir = self.base_out / "work"
        self.metadata_path = self.base_out / "metadata.json"

        self._ensure_dirs()
        self.doc = Document(str(self.input_path))
        self.images_meta = []
        self.pages_meta = []

    def _ensure_dirs(self):
        for p in [self.images_dir, self.pages_dir, self.work_dir]:
            p.mkdir(parents=True, exist_ok=True)

    def _find_rids_in_run(self, run):
        return NS_RID_PATTERN.findall(run._r.xml)

    def extract_images(self):
        parts = self.doc.part.related_parts
        images_meta = []

        # Paragraphs
        for para_idx, para in enumerate(self.doc.paragraphs):
            for run_idx, run in enumerate(para.runs):
                rids = self._find_rids_in_run(run)
                for rid in rids:
                    if rid in parts:
                        part = parts[rid]
                        blob = part.blob
                        ext = part.partname.split('.')[-1]
                        fname = f"img_{para_idx}_{run_idx}_{rid}.{ext}"
                        fpath = self.images_dir / fname
                        with open(fpath, 'wb') as f:
                            f.write(blob)
                        images_meta.append({
                            "rId": rid,
                            "filename": fname,
                            "paragraph_index": para_idx,
                            "run_index": run_idx,
                            "context": para.text[:120]
                        })

        # Tables
        for t_idx, table in enumerate(self.doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for run_idx, run in enumerate(para.runs):
                            rids = self._find_rids_in_run(run)
                            for rid in rids:
                                if rid in parts:
                                    part = parts[rid]
                                    blob = part.blob
                                    ext = part.partname.split('.')[-1]
                                    fname = f"img_table{t_idx}_{r}_{c}_{rid}.{ext}"
                                    fpath = self.images_dir / fname
                                    with open(fpath, 'wb') as f:
                                        f.write(blob)
                                    images_meta.append({
                                        "rId": rid,
                                        "filename": fname,
                                        "table": t_idx,
                                        "row": r,
                                        "col": c,
                                        "context": para.text[:120]
                                    })

        # Orphan images (header/footer/unknown)
        for partname, part in parts.items():
            if 'image' in (part.content_type or ''):
                if any(m['rId'] == partname for m in images_meta):
                    continue
                ext = part.partname.split('.')[-1]
                fname = f"img_orphan_{partname}.{ext}"
                fpath = self.images_dir / fname
                with open(fpath, 'wb') as f:
                    f.write(part.blob)
                images_meta.append({
                    "rId": partname,
                    "filename": fname,
                    "context": "orphan/header/footer/unknown"
                })

        self.images_meta = images_meta

    def build_clean_docx(self):
        new_doc = Document()
        def run_has_image(run):
            return NS_RID_PATTERN.search(run._r.xml) is not None

        for para in self.doc.paragraphs:
            new_p = new_doc.add_paragraph()
            for run in para.runs:
                if run_has_image(run):
                    continue
                new_run = new_p.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                try:
                    new_run.style = run.style
                except Exception:
                    pass

        for table in self.doc.tables:
            new_table = new_doc.add_table(rows=0, cols=len(table.columns))
            for r_idx, row in enumerate(table.rows):
                cells = new_table.add_row().cells
                for c_idx, cell in enumerate(row.cells):
                    texts = []
                    for para in cell.paragraphs:
                        txt_runs = [run.text for run in para.runs if not run_has_image(run)]
                        texts.append(''.join(txt_runs))
                    cells[c_idx].text = '\n'.join(texts)

        self.cleaned_docx_path = self.work_dir / f"{self.input_path.stem}_cleaned.docx"
        new_doc.save(self.cleaned_docx_path)
        return self.cleaned_docx_path

    def convert_to_pdf(self):
        self.pdf_path = self.work_dir / f"{self.input_path.stem}_cleaned.pdf"
        try:
            from docx2pdf import convert
            convert(str(self.cleaned_docx_path), str(self.pdf_path))
            return self.pdf_path
        except Exception:
            soffice = shutil.which('soffice') or shutil.which('libreoffice')
            if not soffice:
                raise RuntimeError("Install docx2pdf or LibreOffice for PDF conversion")
            subprocess.run([soffice, '--headless', '--convert-to', 'pdf',
                            '--outdir', str(self.work_dir),
                            str(self.cleaned_docx_path)], check=True)
            produced = self.work_dir / (self.cleaned_docx_path.stem + '.pdf')
            if produced.exists():
                produced.rename(self.pdf_path)
                return self.pdf_path
            raise RuntimeError("PDF conversion failed")

    def split_pdf(self):
        doc = fitz.open(str(self.pdf_path))
        pages_meta = []
        for i in range(doc.page_count):
            page = doc.load_page(i)
            new_pdf = fitz.open()
            new_pdf.insert_pdf(doc, from_page=i, to_page=i)
            page_fname = self.pages_dir / f"page_{i+1:04d}.pdf"
            new_pdf.save(str(page_fname))
            txt_fname = self.pages_dir / f"page_{i+1:04d}.txt"
            with open(txt_fname, 'w', encoding='utf-8') as f:
                f.write(page.get_text())
            pages_meta.append({
                "page_number": i+1,
                "pdf": str(page_fname.name),
                "textfile": str(txt_fname.name)
            })
        self.pages_meta = pages_meta

    def save_metadata(self):
        metadata = {
            "attempt_hash": self.hash,
            "original_filename": str(self.input_path.name),
            "images": self.images_meta,
            "pages": self.pages_meta
        }
        with open(self.metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

    def run(self):
        print("Extracting images...")
        self.extract_images()
        print(f"{len(self.images_meta)} images extracted to {self.images_dir}")

        print("Building cleaned DOCX...")
        self.build_clean_docx()
        print(f"Cleaned DOCX saved to {self.cleaned_docx_path}")

        print("Converting DOCX to PDF...")
        self.convert_to_pdf()
        print(f"PDF saved to {self.pdf_path}")

        print("Splitting PDF into pages...")
        self.split_pdf()
        print(f"{len(self.pages_meta)} pages saved to {self.pages_dir}")

        print("Saving metadata...")
        self.save_metadata()
        print(f"Metadata saved to {self.metadata_path}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx", type=str, help="Path to input DOCX file")
    parser.add_argument("--outdir", type=str, default="out", help="Output base directory")
    parser.add_argument("--hash", type=str, default=None, help="Optional attempt hash")
    args = parser.parse_args()

    input_path = Path(args.input_docx)
    outdir = Path(args.outdir)

    processor = DocxProcessor(input_path, outdir, args.hash)
    processor.run()

    print("\nAll done. Review pages and metadata before translation.")


if __name__ == "__main__":
    main()
