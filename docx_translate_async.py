#!/usr/bin/env python3
"""
DOCX Translator using Hugging Face Transformers (local, async, progressive save)

- Uses facebook/nllb-200-distilled-600M for high-quality multilingual translation
- Handles large DOCX by splitting into pages/paragraphs
- Preserves line breaks, structure, and inserts images
- Fully local, no API key required
"""

import asyncio
import json
from pathlib import Path
from docx import Document
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch

# Model configuration
MODEL_NAME = "facebook/nllb-200-distilled-600M"
SRC_LANG = "eng_Latn"  # source language code
TGT_LANG = "ukr_Cyrl"  # target language code

# Load tokenizer and model once (GPU if available)
device = 0 if torch.cuda.is_available() else -1
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)
if device != -1:
    model = model.to(device)


class Translator:
    """
    Translator handles text translation using a Seq2Seq model from Hugging Face
    """
    def __init__(self, tokenizer, model, src_lang=SRC_LANG, tgt_lang=TGT_LANG, device=device):
        self.tokenizer = tokenizer
        self.model = model
        self.src_lang = src_lang
        self.tgt_lang = tgt_lang
        self.device = device

    async def translate_page(self, text: str, page_number: int, max_length: int = 1024) -> str:
        """
        Translate a single page/paragraph using NLLB-style language prefix
        """
        # Prefix text with target language code
        text_to_translate = f">>{self.tgt_lang}<< {text}"

        # Encode input text
        inputs = self.tokenizer(text_to_translate, return_tensors="pt", truncation=True, max_length=max_length)
        if self.device != -1:
            inputs = {k: v.to(self.device) for k, v in inputs.items()}

        # Generate translation asynchronously in a thread
        def sync_generate():
            output_tokens = self.model.generate(
                **inputs,
                max_new_tokens=2 * max_length
            )
            return self.tokenizer.batch_decode(output_tokens, skip_special_tokens=True)[0]

        translated = await asyncio.to_thread(sync_generate)
        print(f"=== Page {page_number} translated ===\n{translated[:200]}...\n")
        return translated


class DocxBuilder:
    """
    Handles reading, writing, and updating DOCX files
    """
    def __init__(self, clean_docx_path: Path, mapping_json_path: Path, images_dir: Path):
        self.doc = Document(str(clean_docx_path))
        self.images_dir = images_dir
        with open(mapping_json_path, "r", encoding="utf-8") as f:
            self.mapping = json.load(f)

    def insert_translated_text(self, page_number: int, translated_text: str):
        paragraphs = self.doc.paragraphs
        para_idx = page_number - 1
        if para_idx < len(paragraphs):
            paragraphs[para_idx].text = translated_text
        else:
            self.doc.add_paragraph(translated_text)

    def insert_images(self):
        for img in self.mapping.get("images", []):
            filename = img["filename"]
            p_idx = img.get("paragraph_index")
            if p_idx is not None and p_idx < len(self.doc.paragraphs):
                para = self.doc.paragraphs[p_idx]
                img_path = self.images_dir / filename
                para.add_run().add_picture(str(img_path), width=None)

    def save(self, output_path: Path):
        self.doc.save(output_path)


async def translate_pages_async(pages_dir: Path, translator: Translator):
    """
    Translate all text files in a directory asynchronously (sequentially)
    """
    txt_files = sorted(pages_dir.glob("*.txt"))
    translated_pages = {}
    for txt_file in txt_files:
        page_number = int(txt_file.stem.split("_")[1])
        text = txt_file.read_text(encoding="utf-8")
        translated = await translator.translate_page(text, page_number)
        translated_pages[page_number] = translated
    return translated_pages


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Translate DOCX pages using Hugging Face Transformers locally")
    parser.add_argument("--pages_dir", type=str, required=True)
    parser.add_argument("--clean_docx", type=str, required=True)
    parser.add_argument("--mapping_json", type=str, required=True)
    parser.add_argument("--images_dir", type=str, required=True)
    parser.add_argument("--output_docx", type=str, required=True)
    parser.add_argument("--target_lang", type=str, default="ukr")
    args = parser.parse_args()

    pages_dir = Path(args.pages_dir)
    images_dir = Path(args.images_dir)
    translator = Translator(tokenizer, model, tgt_lang=args.target_lang)

    # Translate pages
    translated_pages = asyncio.run(translate_pages_async(pages_dir, translator))

    # Build translated DOCX
    builder = DocxBuilder(Path(args.clean_docx), Path(args.mapping_json), images_dir)
    for page_number in sorted(translated_pages.keys()):
        builder.insert_translated_text(page_number, translated_pages[page_number])
        builder.save(Path(args.output_docx))
    builder.insert_images()
    builder.save(Path(args.output_docx))
    print(f"Translated DOCX saved to {args.output_docx}")


if __name__ == "__main__":
    main()
